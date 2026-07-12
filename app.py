import io
import re
import streamlit as st
import pdfplumber

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Custom Styling & Theme Configuration for Streamlit ---
st.set_page_config(page_title="KIST Document Generator", page_icon="📄", layout="centered")

st.markdown("""
    <style>
    .stApp { background-color: #f4f8fb; }
    .stButton>button {
        background-color: #4682B4; color: white; border-radius: 6px;
        border: none; padding: 0.5rem 2rem; font-weight: bold;
    }
    .stButton>button:hover { background-color: #356a95; color: white; }
    div[data-testid="stFileUploader"] {
        border: 2px dashed #b0c4de; padding: 20px; border-radius: 10px; background-color: #ffffff;
    }
    </style>
""", unsafe_allow_html=True)

# =====================================================================
# FONT NOTE
# We now generate a real Word (.docx) file instead of a PDF. Word does
# its own complex-script text shaping (reordering Sinhala pre-base
# vowel signs etc. automatically), so we no longer need to bundle a
# font file with this script or manually reorder any characters.
#
# The only requirement is that whoever OPENS the .docx has a Sinhala-
# capable font installed. "Nirmala UI" ships with Windows 10/11 and
# covers Sinhala, so it's used as the default below. If the machine
# that opens the file is older or doesn't have it, change SINHALA_FONT
# to another installed Sinhala font (e.g. "Iskoola Pota", which also
# ships with Windows, or "Noto Sans Sinhala" if that's installed).
# =====================================================================
SINHALA_FONT = "Nirmala UI"


def set_run_font(run, font_name=SINHALA_FONT, size=9, bold=False):
    """Apply a font to a run, including the 'complex script' (w:cs) and
    'east asia' slots — Sinhala is treated as a complex script by Word,
    so w:cs must be set or Word may fall back to a default font for it."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_cell_text(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_table_fixed_layout(table, col_widths_pt):
    """Force Word to respect explicit column widths instead of autofitting."""
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for row in table.rows:
        for idx, width in enumerate(col_widths_pt):
            row.cells[idx].width = Pt(width)
    for idx, width in enumerate(col_widths_pt):
        table.columns[idx].width = Pt(width)


def add_heading_paragraph(doc, text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


# --- Word Document Generation ---
def generate_cash_sheet_docx(invoices, total_amount):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # ==================== PAGE 1 : DAY CASH SHEET ====================
    add_heading_paragraph(doc, "KIST DAY CASH SHEET", size=16, bold=True)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(10)
    meta_run = meta_p.add_run(
        "දිනය : ___________________    මාර්ගය : ___________________    බිල්පත් ගණන : ___________________"
    )
    set_run_font(meta_run, size=10, bold=True)

    headers = [
        "No", "බිල් අංකය", "කඩේ නම", "ප්‍රමාණය", "BNW",
        "කෑන්සල්", "ඇඩ්ජස්ට්", "ඩීස්", "මුදල්", "ණය", "චෙක්", "රිටන්"
    ]
    col_widths = [22, 62, 90, 50, 32, 42, 42, 30, 45, 45, 45, 32]

    total_rows = 1 + len(invoices) + 5 + 1  # header + invoices + 5 blank + ST row
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for c, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[c], h, size=8.5, bold=True)

    # Invoice rows
    r = 1
    idx = 1
    for item in invoices:
        row_cells = table.rows[r].cells
        set_cell_text(row_cells[0], str(idx), size=8)
        set_cell_text(row_cells[1], item['invoice'], size=8)
        set_cell_text(row_cells[2], "", size=8)
        set_cell_text(row_cells[3], f"{item['amount']:.2f}", size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for c in range(4, 12):
            set_cell_text(row_cells[c], "", size=8)
        r += 1
        idx += 1

    # 5 blank spacer rows
    for _ in range(5):
        row_cells = table.rows[r].cells
        set_cell_text(row_cells[0], str(idx), size=8)
        for c in range(1, 12):
            set_cell_text(row_cells[c], "", size=8)
        r += 1
        idx += 1

    # ST / System Sale summary row
    row_cells = table.rows[r].cells
    set_cell_text(row_cells[0], "ST", size=8.5, bold=True)
    set_cell_text(row_cells[1], "System Sale", size=8.5, bold=True)
    set_cell_text(row_cells[2], "", size=8)
    set_cell_text(row_cells[3], f"{total_amount:,.2f}", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for c in range(4, 12):
        set_cell_text(row_cells[c], "", size=8)
    for c in range(12):
        shade_cell(row_cells[c], "EBF2F8")

    set_table_fixed_layout(table, col_widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    sales_labels = [
        "බිස්කට් සෙල් : _______________________",
        "නෙට්ටා සෙල් : _______________________",
        "වතුර සෙල් : _______________________",
        "මුළු සේල් : _______________________",
    ]
    for label in sales_labels:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label)
        set_run_font(run, size=10, bold=True)

    # ==================== PAGE 2 ====================
    doc.add_page_break()
    build_page_two(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_page_two(doc):
    """Builds the second page of the KIST cash sheet:
       1) Credit-bill payments-received table
       2) Cash balancing summary table
       3) Cash-note counting table
       4) Distance-travelled footer line
       Mirrors the layout of the supplied sample Word document.
    """

    # ---- 1) ණය බිල් වලට මුදල් ලැබීම (Credit bill payments received) ----
    add_heading_paragraph(doc, "ණය බිල් වලට මුදල් ලැබීම", size=11, bold=True,
                           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    credit_headers = ["NO", "බිල්පත් දිනය", "කඩේ නම", "ණය", "මුදල් ගෙවූ ප්‍රමාණය", "ඉතිරිය"]
    credit_col_widths = [28, 75, 140, 65, 90, 85]
    credit_table = doc.add_table(rows=11, cols=6)
    credit_table.style = "Table Grid"
    for c, h in enumerate(credit_headers):
        set_cell_text(credit_table.rows[0].cells[c], h, size=8.5, bold=True)
    for i in range(1, 11):
        row_cells = credit_table.rows[i].cells
        set_cell_text(row_cells[0], str(i), size=8)
        for c in range(1, 6):
            set_cell_text(row_cells[c], "", size=8)
    set_table_fixed_layout(credit_table, credit_col_widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- 2) මුදල් පත්‍ර බැලන්ස් කිරීම (Cash balancing) ----
    add_heading_paragraph(doc, "මුදල් පත්‍ර බැලන්ස් කිරීම", size=11, bold=True,
                           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    balance_rows_labels = [
        ("සිස්ටම් සෙල් එක", True),
        ("FOC", False),
        ("මුළු කැන්සල්", False),
        ("ශේෂය (1)", True),
        ("මුළු වට්ටම්", False),
        ("මුළු ඇඩ්ජස්ට්", False),
        ("මුළු රිටන්", False),
        ("ශේෂය (2)", True),
        ("මුළු මුදල්", False),
        ("මුළු ණය", False),
        ("මුළු චෙක්පත්", False),
        ("ශේෂය (3)", True),
        ("මුදල් ශේෂය", True),
        ("මුළු දින මුදල්", False),
        ("ලැබුණු මුළු ණය", False),
        ("මුළු වියදම්", False),
        ("බැංකුගත වටිනාකම", False),
    ]
    balance_col_widths = [300, 183]
    balance_table = doc.add_table(rows=len(balance_rows_labels), cols=2)
    balance_table.style = "Table Grid"
    for i, (label, is_bold) in enumerate(balance_rows_labels):
        row_cells = balance_table.rows[i].cells
        set_cell_text(row_cells[0], label, size=8, bold=is_bold, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row_cells[1], "", size=8)
    set_table_fixed_layout(balance_table, balance_col_widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- 3) මුදල් ගණනය කිරීම (Cash note counting) ----
    add_heading_paragraph(doc, "මුදල් ගණනය කිරීම", size=11, bold=True,
                           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    note_headers = ["මුදල් නෝට්ටු", "වටිනාකම"]
    note_values = ["20", "50", "100", "500", "1000", "2000", "5000", "coins", "Total"]
    note_col_widths = [240, 240]
    note_table = doc.add_table(rows=1 + len(note_values), cols=2)
    note_table.style = "Table Grid"
    for c, h in enumerate(note_headers):
        set_cell_text(note_table.rows[0].cells[c], h, size=8.5, bold=True)
    for i, note in enumerate(note_values, start=1):
        row_cells = note_table.rows[i].cells
        is_total = (note == "Total")
        set_cell_text(row_cells[0], note, size=8, bold=is_total)
        set_cell_text(row_cells[1], "", size=8)
        if is_total:
            shade_cell(row_cells[0], "EBF2F8")
            shade_cell(row_cells[1], "EBF2F8")
    set_table_fixed_layout(note_table, note_col_widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- 4) Footer: distance travelled ----
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        "ගමන් කළ දුර : ___________    KM : ___________    OOT : ___________"
    )
    set_run_font(footer_run, size=10, bold=True)


# --- Stateful Line-by-Line Multi-Format Text Extraction Engine ---
def parse_pdf_file(uploaded_file):
    invoices = []
    total_amount = 0.0
    invoice_seen = set()

    composite_invoice_pattern = re.compile(r'\b\d{2}[A-Z]{3}_\d{4,}[\w\d_]*\b')
    standard_invoice_pattern = re.compile(r'\b(IN|TI)\d{5,7}\b')
    isolated_id_pattern = re.compile(r'^\b\d{3,4}\b$')
    amount_pattern = re.compile(r'\d[\d,]*\.\d{2}')

    with pdfplumber.open(uploaded_file) as pdf:
        active_prefix = None
        active_sub_id = None

        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue

            lines = text.splitlines()
            for line in lines:
                line_str = line.strip()
                if not line_str:
                    continue

                if line_str.lower().startswith("total") or "grand total" in line_str.lower():
                    amts = amount_pattern.findall(line_str)
                    if amts:
                        try:
                            total_amount = float(amts[-1].replace(",", ""))
                        except ValueError:
                            pass
                    continue

                comp_match = composite_invoice_pattern.search(line_str)
                if comp_match:
                    active_prefix = comp_match.group()
                    line_remainder = line_str.replace(active_prefix, "").strip()
                    remainder_numbers = re.findall(r'\b\d{3,4}\b', line_remainder)
                    if remainder_numbers:
                        active_sub_id = remainder_numbers[0]
                    continue

                std_match = standard_invoice_pattern.search(line_str)
                if std_match:
                    active_prefix = std_match.group()
                    active_sub_id = None

                iso_match = isolated_id_pattern.search(line_str)
                if iso_match and active_prefix and not active_prefix.startswith(("IN", "TI")):
                    active_sub_id = iso_match.group()
                    continue

                amts = amount_pattern.findall(line_str)
                if amts and active_prefix:
                    if "net amt" in line_str.lower() or "mrp" in line_str.lower():
                        continue

                    if active_prefix.startswith(("IN", "TI")):
                        invoice_identity = active_prefix
                    elif active_sub_id:
                        invoice_identity = active_sub_id
                    else:
                        digits_found = re.findall(r'\d+', active_prefix)
                        invoice_identity = digits_found[-1] if digits_found else active_prefix

                    try:
                        amt_val = float(amts[-1].replace(",", ""))

                        if len(invoice_identity) >= 2:
                            slice_length = 4
                            final_invoice_slice = invoice_identity[-slice_length:]

                            while final_invoice_slice in invoice_seen and slice_length < len(invoice_identity):
                                slice_length += 1
                                final_invoice_slice = invoice_identity[-slice_length:]

                            invoices.append({
                                "invoice": final_invoice_slice,
                                "amount": amt_val
                            })
                            invoice_seen.add(final_invoice_slice)

                        active_sub_id = None
                        if active_prefix.startswith(("IN", "TI")):
                            active_prefix = None
                    except ValueError:
                        pass

    if total_amount == 0.0 and invoices:
        total_amount = sum(i["amount"] for i in invoices)

    return invoices, total_amount


# --- Main App Interface ---
st.html("""
    <div style="margin-bottom: 1.5rem;">
        <h1 style="color: #000000 !important; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin-bottom: 0.5rem; font-size: 2.25rem; font-weight: 700;">
            📄 KIST Sheet Custom Generation Engine
        </h1>
        <p style="color: #000000 !important; font-size: 1rem; font-weight: normal; margin: 0;">
            Upload the picklist PDF file below to instantly extract data and generate the structured layout.
        </p>
    </div>
""")

uploaded_file = st.file_uploader("Select input Picklist PDF file", type=["pdf"])

if uploaded_file is not None:
    invoices, total_amt = parse_pdf_file(uploaded_file)

    if invoices:
        st.success(f"Successfully processed {len(invoices)} invoices records from PDF. Identified System Sale Value: LKR {total_amt:,.2f}")

        with st.spinner("Compiling structural printable layout..."):
            docx_data = generate_cash_sheet_docx(invoices, total_amt)

        st.download_button(
            label="📥 Download Tailored Cash Sheet (Word)",
            data=docx_data,
            file_name="KIST_Day_Cash_Sheet.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Could not parse invoices structures. Please verify that the PDF contains valid invoice entries.")
